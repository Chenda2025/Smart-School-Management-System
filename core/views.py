# core/views.py
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from users.models import CustomUser
from .models import Student, ClassRoom, Section, Teacher,Subject,Attendance,Score
from datetime import date

@login_required
def add_student(request):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can add students!")
        return redirect('all_users_list')

    # Users who are not already students
    existing_student_ids = Student.objects.values_list('user_id', flat=True)
    available_users = CustomUser.objects.exclude(id__in=existing_student_ids).exclude(role__in=['teacher', 'admin'])

    if request.method == 'POST':
        user_id = request.POST.get('user')
        roll_number = request.POST.get('roll_number')
        classroom_id = request.POST.get('classroom')
        section_id = request.POST.get('section')

        if Student.objects.filter(roll_number=roll_number).exists():
            messages.error(request, "Roll number already exists!")
        else:
            user = CustomUser.objects.get(id=user_id)
            Student.objects.create(
                user=user,
                roll_number=roll_number,
                classroom=ClassRoom.objects.get(id=classroom_id),
                section=Section.objects.get(id=section_id),
            )
            messages.success(request, f"Student {user.get_full_name()} added!")
            return redirect('all_users_list')

    context = {
        'available_users': available_users,
        'classrooms': ClassRoom.objects.all(),
        'sections': Section.objects.all(),
    }
    return render(request, 'core/add_student.html', context)

# core/views.py  ← ONLY THIS PART (replace everything from @login_required down)

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from users.models import CustomUser
from .models import Subject      # ← ONLY THESE TWO NEEDED


# core/views.py — FINAL FINAL FINAL VERSION

@login_required
def add_teacher(request):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can add teachers!")
        return redirect('all_users_list')

    available_users = CustomUser.objects.exclude(
        teacher_profile__isnull=False
    ).exclude(role='admin').order_by('first_name', 'last_name')

    all_subjects = Subject.objects.select_related('classroom').all().order_by('classroom__name', 'name')

    if request.method == 'POST':
        user_id = request.POST.get('user')
        employee_id = request.POST.get('employee_id', '').strip()
        subject_ids = request.POST.getlist('subjects')

        if not user_id or not employee_id or not subject_ids:
            messages.error(request, "All fields are required!")
        elif CustomUser.objects.filter(employee_id=employee_id).exclude(pk=user_id).exists():
            messages.error(request, "Employee ID already taken!")
        else:
            user = CustomUser.objects.get(id=user_id)

            # ONLY THESE 4 LINES — NO Teacher.objects.create() ANYMORE
            user.employee_id = employee_id
            user.role = 'teacher'
            user.subjects_taught.set(Subject.objects.filter(id__in=subject_ids))
            user.save()

            messages.success(request, f"Teacher '{user.get_full_name()} added successfully!")
            return redirect('all_users_list')

    context = {
        'available_users': available_users,
        'subject': all_subjects,
    }
    return render(request, 'core/add_teacher.html', context)


from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.hashers import make_password
from django.contrib.auth.decorators import login_required
import pandas as pd
from io import BytesIO
from core.models import Student, ClassRoom, Section
from users.models import CustomUser

@login_required
def bulk_import(request):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can use bulk import!")
        return redirect('all_users_list')

    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        try:
            df = pd.read_excel(file) if file.name.endswith('.xlsx') else pd.read_csv(file)
            request.session['bulk_data'] = df.to_dict('records')
            return redirect('bulk_import_preview')
        except Exception as e:
            messages.error(request, f"Error reading file: {e}")

    return render(request, 'core/bulk_import.html')

@login_required
def bulk_import_preview(request):
    if not request.user.is_superuser:
        return redirect('all_users_list')

    data = request.session.get('bulk_data', [])
    if not data:
        messages.error(request, "No data to import!")
        return redirect('bulk_import')

    if request.method == 'POST':
        success = 0
        errors = []
        for i, row in enumerate(data):
            try:
                username = str(row.get('username') or row.get('Username') or '').strip()
                if not username:
                    errors.append(f"Row {i+2}: Username missing")
                    continue
                if CustomUser.objects.filter(username=username).exists():
                    errors.append(f"Row {i+2}: Username {username} already exists")
                    continue

                # Create user
                user = CustomUser.objects.create(
                    username=username,
                    first_name=row.get('first_name') or row.get('First Name') or '',
                    last_name=row.get('last_name') or row.get('Last Name') or '',
                    email=row.get('email') or row.get('Email') or '',
                    role='student',
                    password=make_password("123456"),  # default password
                    is_active=True
                )

                # Create student
                roll = str(row.get('roll_number') or row.get('Roll Number') or '')
                class_name = str(row.get('class') or row.get('Class') or '')
                section_name = str(row.get('section') or row.get('Section') or '')

                classroom = ClassRoom.objects.get(name__iexact=class_name)
                section = Section.objects.get(name__iexact=section_name, classroom=classroom)

                Student.objects.create(
                    user=user,
                    roll_number=roll,
                    classroom=classroom,
                    section=section
                )
                success += 1
            except Exception as e:
                errors.append(f"Row {i+2}: {str(e)}")

        if errors:
            messages.warning(request, f"{success} imported. Errors: {len(errors)}")
            for err in errors[:10]:
                messages.error(request, err)
        else:
            messages.success(request, f"Successfully imported {success} students!")
        del request.session['bulk_data']
        return redirect('all_users_list')

    context = {'data': data[:20], 'total': len(data)}  # Show preview
    return render(request, 'core/bulk_import_preview.html', context)

# core/views.py — ADD THESE TWO FUNCTIONS
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import ClassRoom, Section

@login_required
def classes_list(request):
    if request.method == 'POST':
        action = request.POST.get('action')

        # ADD CLASS
        if action == 'add_class':
            name = request.POST.get('class_name', '').strip()
            if name:
                if ClassRoom.objects.filter(name__iexact=name).exists():
                    messages.error(request, f"Class '{name}' already exists!")
                else:
                    ClassRoom.objects.create(name=name)
                    messages.success(request, f"Class '{name}' added!")
            else:
                messages.error(request, "Class name cannot be empty!")

        # ADD SECTION
        elif action == 'add_section':
            name = request.POST.get('section_name', '').strip()
            classroom_id = request.POST.get('classroom')
            if name and classroom_id:
                classroom = get_object_or_404(ClassRoom, id=classroom_id)
                if Section.objects.filter(name__iexact=name, classroom=classroom).exists():
                    messages.error(request, "Section already exists in this class!")
                else:
                    Section.objects.create(name=name, classroom=classroom)
                    messages.success(request, f"Section '{name}' added to {classroom}!")
            else:
                messages.error(request, "Fill all fields!")

        # EDIT CLASS
        elif action == 'edit_class':
            class_id = request.POST.get('class_id')
            new_name = request.POST.get('new_name', '').strip()
            if new_name and class_id:
                classroom = get_object_or_404(ClassRoom, id=class_id)
                if ClassRoom.objects.filter(name__iexact=new_name).exclude(id=class_id).exists():
                    messages.error(request, "Class name already taken!")
                else:
                    classroom.name = new_name
                    classroom.save()
                    messages.success(request, "Class updated!")

        # DELETE CLASS
        elif action == 'delete_class':
            class_id = request.POST.get('class_id')
            classroom = get_object_or_404(ClassRoom, id=class_id)
            name = classroom.name
            classroom.delete()
            messages.success(request, f"Class '{name}' deleted!")

        # DELETE SECTION
        elif action == 'delete_section':
            section_id = request.POST.get('section_id')
            section = get_object_or_404(Section, id=section_id)
            name = section.name
            classroom = section.classroom
            section.delete()
            messages.success(request, f"Section '{name}' removed from {classroom}!")

        return redirect('classes_list')

    classrooms = ClassRoom.objects.prefetch_related('sections').all().order_by('name')
    return render(request, 'core/classes_list.html', {'classrooms': classrooms})

@login_required
def subjects_list(request):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can manage subjects!")
        return redirect('all_users_list')

    if request.method == 'POST':
        action = request.POST.get('action')

        # ADD SUBJECT
        if action == 'add_subject':
            name = request.POST.get('name', '').strip()
            class_id = request.POST.get('classroom')
            teacher_id = request.POST.get('teacher')

            if not all([name, class_id, teacher_id]):
                messages.error(request, "All fields required!")
            elif Subject.objects.filter(name__iexact=name, classroom_id=class_id).exists():
                messages.error(request, f"'{name}' already exists in this class!")
            else:
                Subject.objects.create(
                    name=name,
                    classroom=ClassRoom.objects.get(id=class_id),
                    teacher=CustomUser.objects.get(id=teacher_id)
                )
                messages.success(request, f"Subject '{name}' added!")

       # core/views.py — FINAL EDIT + DUPLICATE PROTECTION
        elif action == 'edit_subject':
            subject_id = request.POST.get('subject_id')
            name = request.POST.get('name', '').strip()
            class_id = request.POST.get('classroom')
            teacher_id = request.POST.get('teacher')

            subject = get_object_or_404(Subject, id=subject_id)

            # Prevent duplicate: same subject name in same class
            if Subject.objects.filter(name__iexact=name, classroom_id=class_id).exclude(id=subject_id).exists():
                messages.error(request, f"'{name}' already exists in this class!")
            else:
                subject.name = name
                subject.classroom = ClassRoom.objects.get(id=class_id)
                subject.teacher = CustomUser.objects.get(id=teacher_id) if teacher_id else None
                subject.save()
                messages.success(request, "Subject updated successfully!")

        # DELETE SUBJECT
        elif action == 'delete_subject':
            subject = get_object_or_404(Subject, id=request.POST.get('subject_id'))
            name = subject.name
            subject.delete()
            messages.success(request, f"Subject '{name}' deleted!")

        return redirect('subjects_list')

    subjects = Subject.objects.select_related('classroom', 'teacher').all().order_by('classroom__name', 'name')
    context = {
        'subjects': subjects,
        'classrooms': ClassRoom.objects.all(),
        'teachers': CustomUser.objects.filter(teacher_profile__isnull=False),
    }
    return render(request, 'core/subjects_list.html', context)

  
# core/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from core.models import Student, ClassRoom, Subject, Attendance


@login_required
def mark_attendance(request):
    # Only teachers & admin can mark attendance
    if not request.user.is_superuser and not hasattr(request.user, 'teacher_profile'):
        messages.error(request, "Only teachers and admin can mark attendance!")
        return redirect('staff_dashboard')

    today = timezone.now().date()
    selected_grade = request.GET.get('grade')

    # Base querysets
    classrooms = ClassRoom.objects.all()
    selected_classroom = None
    students = Student.objects.none()
    subjects = Subject.objects.none()

    # Filter by selected grade
    if selected_grade:
        try:
            selected_classroom = ClassRoom.objects.get(id=selected_grade)
            students = Student.objects.filter(classroom=selected_classroom).select_related('user')
            subjects = Subject.objects.filter(classroom=selected_classroom)
        except ClassRoom.DoesNotExist:
            messages.error(request, "Invalid class selected!")
            selected_grade = None

    # TODAY'S ATTENDANCE LIST — FOR MODAL
    today_attendance = Attendance.objects.filter(
        date=today
    ).select_related('student', 'student__user', 'subject', 'student__classroom').order_by('-recorded_at')

    # POST — SAVE ATTENDANCE
    if request.method == 'POST':
        student_id = request.POST.get('student')
        subject_id = request.POST.get('subject')
        status = request.POST.get('status')

        if not all([student_id, subject_id, status]):
            messages.error(request, "Please fill all fields!")
        else:
            try:
                student = Student.objects.get(id=student_id)
                subject = Subject.objects.get(id=subject_id)

                # Prevent double entry today
                if Attendance.objects.filter(student=student, subject=subject, date=today).exists():
                    messages.error(request, f"Attendance already marked today for {student} in {subject}!")
                else:
                    Attendance.objects.create(
                        student=student,
                        subject=subject,
                        status=status,
                        marked_by=request.user
                    )

                    status_display = dict(Attendance.STATUS_CHOICES).get(status, status).title()
                    messages.success(request, f"{student.user.get_full_name()} → {status_display}")

                    # Auto warning system
                    absences = Attendance.objects.filter(
                        student=student,
                        subject=subject,
                        status='absent'
                    ).count()
                    if absences == 2:
                        messages.warning(request, f"Warning: {student} has 2 absences in {subject} — ONE MORE = SCORE 0!")
                    elif absences >= 3:
                        messages.error(request, f"CRITICAL: {student} has {absences} absences → Score will be 0!")

            except Exception as e:
                messages.error(request, "Error saving attendance!")

        return redirect('mark_attendance')

    context = {
        'classrooms': classrooms,
        'students': students.order_by('user__first_name'),
        'subjects': subjects,
        'selected_grade': selected_grade,
        'selected_classroom': selected_classroom,
        'today': today,
        'total_marked_today': Attendance.objects.filter(date=today, marked_by=request.user).count(),
        'today_attendance': today_attendance,
    }
    return render(request, 'core/mark_attendance.html', context)

@login_required
def edit_attendance(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can edit!")
        return redirect('mark_attendance')
    
    record = get_object_or_404(Attendance, pk=pk)
    if request.method == 'POST':
        record.status = request.POST.get('status')
        record.save()
        messages.success(request, "Attendance updated!")
        return redirect('mark_attendance')
    
    context = {'record': record}
    return render(request, 'core/edit_attendance.html', context)

@login_required
def delete_attendance(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can delete!")
        return redirect('mark_attendance')
    
    record = get_object_or_404(Attendance, pk=pk)
    if request.method == 'POST':
        record.delete()
        messages.success(request, "Attendance deleted!")
        return redirect('mark_attendance')
    
    return redirect('mark_attendance')

import json
from decimal import Decimal

from django.contrib.auth.decorators import login_required
from django.db.models import Avg, Count, FloatField
from django.shortcuts import render
from django.core.serializers.json import DjangoJSONEncoder

@login_required
def legacy_scores(request):
    if not request.user.is_superuser and not hasattr(request.user, 'teacher_profile'):
        messages.error(request, "Access restricted to teachers and admins.")
        return redirect('staff_dashboard')

    # Filters
    selected_class_id = request.GET.get('class')
    selected_exam_type = request.GET.get('exam_type', '').lower() or None
    selected_year = request.GET.get('year')

    # Base querysets
    scores_qs = Score.objects.select_related(
        'student', 'subject', 'student__user', 'student__classroom'
    ).order_by('-recorded_at')

    # Apply filters SAFELY using field names, not _id
    if selected_class_id:
        scores_qs = scores_qs.filter(student__classroom__id=selected_class_id)
        # Or even safer: scores_qs = scores_qs.filter(student__classroom=selected_class_id)

    if selected_exam_type:
        scores_qs = scores_qs.filter(exam_type=selected_exam_type)

    if selected_year:
        scores_qs = scores_qs.filter(recorded_at__year=selected_year)

    # Available filters
    classrooms = ClassRoom.objects.all().order_by('name')
    available_years = Score.objects.dates('recorded_at', 'year').values_list('recorded_at__year', flat=True).distinct()

    # === STATS ===
    total_scores = scores_qs.count()
    avg_score = scores_qs.aggregate(avg=Avg('score'))['avg']
    avg_score = round(avg_score or 0.0, 1) if avg_score else 0.0

    # Grade distribution
    grade_dist = scores_qs.values('grade').annotate(count=Count('id')).order_by('grade')

    # === SUBJECT PERFORMANCE ===
    subject_perf_qs = Subject.objects.annotate(
        avg=Avg('scores__score')
    ).filter(avg__isnull=False)

    if selected_class_id:
        subject_perf_qs = subject_perf_qs.filter(scores__student__classroom__id=selected_class_id)
    if selected_exam_type:
        subject_perf_qs = subject_perf_qs.filter(scores__exam_type=selected_exam_type)
    if selected_year:
        subject_perf_qs = subject_perf_qs.filter(scores__recorded_at__year=selected_year)

    subject_perf = list(subject_perf_qs.values('name', 'avg').order_by('-avg')[:15])
    for item in subject_perf:
        item['avg'] = round(item['avg'] or 0, 1)

    # === ALL-TIME TOP 10 ===
    historical_top = list(
        Student.objects.annotate(
            overall_avg=Avg('scores__score')
        )
        .filter(overall_avg__isnull=False)
        .select_related('user', 'classroom')
        .order_by('-overall_avg')[:10]
        .values('user__first_name', 'user__last_name', 'roll_number', 'classroom__name', 'overall_avg')
    )
    for s in historical_top:
        s['overall_avg'] = round(s['overall_avg'] or 0, 1)

    # === DETAILED SCORES TABLE ===
    detailed_scores = list(
        scores_qs.values(
            'student__user__first_name', 'student__user__last_name',
            'student__roll_number', 'student__classroom__name',
            'subject__name', 'exam_type', 'score', 'grade', 'recorded_at'
        )[:500]
    )

    for row in detailed_scores:
        row['full_name'] = f"{row['student__user__first_name']} {row['student__user__last_name']}".strip()
        row['recorded_date'] = row['recorded_at'].strftime('%b %d, %Y') if row['recorded_at'] else '-'
        row['score_display'] = f"{row['score']:.1f}" if row['score'] is not None else '—'

    context = {
        'classrooms': classrooms,
        'available_years': sorted(available_years, reverse=True),
        'exam_types': [choice for choice in Score.EXAM_TYPES],

        'selected_class_id': selected_class_id,
        'selected_exam_type': selected_exam_type.capitalize() if selected_exam_type else '',
        'selected_year': selected_year,

        'total_scores': total_scores,
        'avg_score': avg_score,

        'grade_labels': json.dumps([item['grade'] or 'N/A' for item in grade_dist]),
        'grade_values': json.dumps([item['count'] for item in grade_dist]),

        'subject_perf': json.dumps(subject_perf, cls=DjangoJSONEncoder),
        'historical_top': json.dumps(historical_top, cls=DjangoJSONEncoder),
        'detailed_scores': detailed_scores,
    }

    return render(request, 'core/legacy_scores.html', context)

# core/views.py
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from datetime import timedelta
from django.core.paginator import Paginator
from core.models import Student, ClassRoom, Subject, Attendance


@login_required
def student_info(request):
    # GET FILTERS
    selected_grade = request.GET.get('grade')
    selected_subject = request.GET.get('subject')
    period = request.GET.get('period', '')

    # CLEAN 'None' FROM URL
    if selected_grade in ('', 'None', None):
        selected_grade = None
    if selected_subject in ('', 'None', None):
        selected_subject = None

    # BASE QUERY — ATTENDANCE RECORDS
    attendance_qs = Attendance.objects.select_related(
        'student', 'student__user', 'subject', 'student__classroom'
    ).order_by('-date', '-recorded_at')

    # FILTER BY CLASS (GRADE)
    if selected_grade:
        attendance_qs = attendance_qs.filter(student__classroom_id=selected_grade)

    # FILTER BY SUBJECT
    if selected_subject:
        attendance_qs = attendance_qs.filter(subject_id=selected_subject)

    # FILTER BY TIME PERIOD
    today = timezone.now().date()
    if period == 'last_week':
        attendance_qs = attendance_qs.filter(date__gte=today - timedelta(days=7))
    elif period == '1_month':
        attendance_qs = attendance_qs.filter(date__gte=today - timedelta(days=30))
    elif period == '1_year':
        attendance_qs = attendance_qs.filter(date__gte=today - timedelta(days=365))


    # PAGINATION — THIS IS THE FINAL FIX
    paginator = Paginator(attendance_qs, 15)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)  # ← THIS LINE FIXES THE ERROR

    # DYNAMIC SUBJECTS FOR DROPDOWN — ONLY FROM SELECTED GRADE
    subjects_for_dropdown = Subject.objects.select_related('classroom')
    if selected_grade:
        subjects_for_dropdown = subjects_for_dropdown.filter(classroom_id=selected_grade)
    else:
        subjects_for_dropdown = subjects_for_dropdown.all()

    context = {
        'page_obj': page_obj,
        'classrooms': ClassRoom.objects.all(),
        'subjects': subjects_for_dropdown,           # ← ONLY SUBJECTS BY GRADE
        'selected_grade': selected_grade,
        'selected_subject': selected_subject,
        'period': period,
        'today': today,
    }
    return render(request, 'core/student_info.html', context)

# core/views.py — FINAL WORKING EXPORTS (NO WEASYPRINT!)

from django.http import HttpResponse
from django.utils import timezone
import xlsxwriter
from io import BytesIO
from docx import Document
from docx.shared import Inches

# EXCEL EXPORT — 100% WORKING
@login_required
def export_attendance_excel(request):
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet('Attendance')

    # Header
    title_format = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center'})
    worksheet.merge_range('A1:I1', 'ELITE SCHOOL - ATTENDANCE REPORT', title_format)
    worksheet.write('A2', f"Generated: {timezone.now().strftime('%Y-%m-%d %H:%M')}")

    # Data
    records = Attendance.objects.select_related('student', 'student__user', 'subject', 'student__classroom')
    if request.GET.get('grade'):
        records = records.filter(student__classroom_id=request.GET['grade'])
    if request.GET.get('subject'):
        records = records.filter(subject_id=request.GET['subject'])

    headers = ['ID', 'Name', 'Roll No', 'Class', 'Subject', 'Date', 'Time', 'Status']
    row = 3
    worksheet.write_row(row, 0, headers)
    row += 1

    for r in records:
        worksheet.write_row(row, 0, [
            r.id,
            r.student.user.get_full_name(),
            r.student.roll_number,
            r.student.classroom.name,
            r.subject.name,
            r.date.strftime('%Y-%m-%d'),
            r.recorded_at.strftime('%H:%M'),
            r.get_status_display()
        ])
        row += 1

    workbook.close()
    output.seek(0)

    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=attendance_elite_school.xlsx'
    return response

# WORD EXPORT — 100% WORKING
@login_required
def export_attendance_word(request):
    doc = Document()
    doc.add_heading('ELITE SCHOOL - ATTENDANCE REPORT', 0)
    doc.add_paragraph(f"Generated on: {timezone.now().strftime('%Y-%m-%d %H:%M')}")

    table = doc.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    headers = ['ID', 'Name', 'Roll No', 'Class', 'Subject', 'Date', 'Time', 'Status']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    records = Attendance.objects.select_related('student', 'student__user', 'subject', 'student__classroom')
    if request.GET.get('grade'):
        records = records.filter(student__classroom_id=request.GET['grade'])
    if request.GET.get('subject'):
        records = records.filter(subject_id=request.GET['subject'])

    for r in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.id)
        row_cells[1].text = r.student.user.get_full_name()
        row_cells[2].text = r.student.roll_number
        row_cells[3].text = r.student.classroom.name
        row_cells[4].text = r.subject.name
        row_cells[5].text = r.date.strftime('%Y-%m-%d')
        row_cells[6].text = r.recorded_at.strftime('%H:%M')
        row_cells[7].text = r.get_status_display()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=attendance_elite_school.docx'
    return response

# core/views.py
# core/views.py — TOP OF FILE
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from core.models import Student, ClassRoom, Subject, Attendance, QRSession  # ← QRSession ADDED HERE
@login_required
def qr_attendance_home(request):
    if request.user.is_superuser:
        # SuperStaff — Can create QR
        classrooms = ClassRoom.objects.all()
        subjects = Subject.objects.all()
        active_session = QRSession.objects.filter(is_active=True).first()
        return render(request, 'core/qr_create.html', {
            'classrooms': classrooms,
            'subjects': subjects,
            'active_session': active_session,
        })
    else:
        # Teacher or Student — Only scan
        return render(request, 'core/qr_scan.html')


# # core/views.py
# @login_required
# def qr_attendance_home(request):
#     if request.user.is_superuser:
#         # SuperStaff — Can create QR
#         classrooms = ClassRoom.objects.all()
#         subjects = Subject.objects.all()
#         active_session = QRSession.objects.filter(is_active=True).first()
#         return render(request, 'core/qr_create.html', {
#             'classrooms': classrooms,
#             'subjects': subjects,
#             'active_session': active_session,
#         })
#     else:
#         # Teacher or Student — Only scan
#         return render(request, 'core/qr_scan.html')


@login_required
def create_qr_session(request):
    if not request.user.is_superuser:
        messages.error(request, "Only admin can create QR!")
        return redirect('qr_attendance_home')

    if request.method == 'POST':
        classroom_id = request.POST.get('classroom')
        subject_id = request.POST.get('subject')

        # Deactivate old sessions
        QRSession.objects.filter(is_active=True).update(is_active=False)

        # Create new
        session = QRSession.objects.create(
            created_by=request.user,
            classroom_id=classroom_id,
            subject_id=subject_id or None,
            expires_at=timezone.now() + timedelta(minutes=2)
        )
        messages.success(request, "QR Session Created! Valid for 2 minutes.")
        return redirect('qr_attendance_home')

    return redirect('qr_attendance_home')


def qr_scan_view(request, token):
    try:
        session = QRSession.objects.get(token=token, is_active=True, expires_at__gt=timezone.now())
    except QRSession.DoesNotExist:
        return render(request, 'core/qr_invalid.html')

    # Mark attendance
    Attendance.objects.update_or_create(
        student=request.user.student,
        subject=session.subject,
        date=timezone.now().date(),
        defaults={'status': 'present', 'marked_by': session.created_by}
    )
    return render(request, 'core/qr_success.html', {'session': session})


# from django.shortcuts import render, redirect, get_object_or_404
# from django.contrib.auth.decorators import login_required
# from django.contrib import messages
# from django.db.models import Q
# from core.models import Student, ClassRoom, Subject, Score, Attendance

from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.urls import reverse
from urllib.parse import urlencode
from decimal import Decimal

@login_required
def enter_scores(request):
    # Permission check: Only teachers or superusers
    if not request.user.is_superuser and not hasattr(request.user, 'teacher_profile'):
        messages.error(request, "Only teachers and admins can enter scores!")
        return redirect('staff_dashboard')

    # Safely get selected class and exam_type
    selected_class_id = request.GET.get('class') or request.POST.get('classroom')
    if selected_class_id in ('', 'None', None):
        selected_class_id = None

    # Get exam_type from POST (save) or GET (filter), default to 'midterm'
    exam_type = (request.POST.get('exam_type') or request.GET.get('exam_type', 'midterm')).lower()

    # Validate exam_type against model choices
    valid_exam_types = {choice[0].lower(): choice[1] for choice in Score.EXAM_TYPES}
    if exam_type not in valid_exam_types:
        exam_type = 'midterm'

    # Base querysets
    classrooms = ClassRoom.objects.all()
    selected_classroom = None
    subjects_in_class = Subject.objects.none()
    all_students_in_class = Student.objects.none()
    students_without_scores = Student.objects.none()
    student_scores = {}  # {student_id: {subject_id: score}}

    if selected_class_id:
        try:
            selected_classroom = ClassRoom.objects.get(id=selected_class_id)
            subjects_in_class = Subject.objects.filter(classroom=selected_classroom).order_by('name')

            # All students in class
            all_students_in_class = Student.objects.filter(
                classroom=selected_classroom
            ).select_related('user').order_by('roll_number')

            # Load existing scores for this exam_type
            existing_scores = Score.objects.filter(
                student__classroom=selected_classroom,
                exam_type=exam_type
            ).select_related('student', 'subject')

            # Build student_scores dict for template
            for score in existing_scores:
                student_scores.setdefault(score.student_id, {})[score.subject_id] = score.score

            # Students who haven't received any score yet
            scored_student_ids = existing_scores.values_list('student_id', flat=True).distinct()
            students_without_scores = all_students_in_class.exclude(id__in=scored_student_ids)

        except ClassRoom.DoesNotExist:
            messages.error(request, "Selected class does not exist.")
            selected_class_id = None

    # POST: Save all scores
    if request.method == 'POST' and selected_class_id:
        saved_count = 0
        warning_shown = set()  # Prevent duplicate absence warnings

        with transaction.atomic():
            for key, value in request.POST.items():
                if not key.startswith('score_') or not value.strip():
                    continue

                try:
                    _, student_id, subject_id = key.split('_')
                    raw_score = Decimal(value.strip())

                    # Basic validation
                    if not (0 <= raw_score <= 100):
                        continue

                    student = Student.objects.get(id=student_id, classroom=selected_classroom)
                    subject = Subject.objects.get(id=subject_id, classroom=selected_classroom)

                    # Check absences → auto-set to 0.00 if ≥3
                    absences = Attendance.objects.filter(
                        student=student,
                        subject=subject,
                        status='absent'
                    ).count()

                    final_score = Decimal('0.00') if absences >= 3 else raw_score

                    if absences >= 3 and student_id not in warning_shown:
                        messages.warning(
                            request,
                            f"{student.user.get_full_name()} has {absences} absences in {subject.name} → score set to 0.00"
                        )
                        warning_shown.add(student_id)

                    # Save or update score — model handles grade, validation, recorded_by, timestamp
                    Score.objects.update_or_create(
                        student=student,
                        subject=subject,
                        exam_type=exam_type,
                        defaults={
                            'score': final_score,
                            'recorded_by': request.user,
                        }
                    )
                    saved_count += 1

                except (ValueError, Student.DoesNotExist, Subject.DoesNotExist):
                    continue  # Skip invalid/malformed inputs

        # Success messages
        if saved_count > 0:
            messages.success(request, f"Successfully saved {saved_count} score(s)! Grades auto-calculated.")
        else:
            messages.info(request, "No valid scores were submitted.")

        # FIXED REDIRECT: Properly preserve query parameters using reverse()
        query_params = {
            'class': selected_class_id,
            'exam_type': valid_exam_types[exam_type]  # Capitalized display: Quiz, Midterm, etc.
        }
        redirect_url = reverse('enter_scores') + '?' + urlencode(query_params)
        return redirect(redirect_url)

    # Template context
    context = {
        'classrooms': classrooms,
        'selected_classroom': selected_classroom,
        'selected_grade': selected_class_id,
        'subjects_in_class': subjects_in_class,
        'subjects': subjects_in_class,  # backward compatibility
        'all_students_in_class': all_students_in_class,
        'students_in_class': students_without_scores,
        'student_scores': student_scores,
        'exam_type': valid_exam_types.get(exam_type, 'Midterm'),  # Display name
        'exam_type_raw': exam_type,  # Raw value for hidden input
    }

    return render(request, 'core/enter_scores.html', context)

from django.db.models import Avg, Count, F
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
import json

from django.db.models import Avg
import json
from decimal import Decimal

class DecimalEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, Decimal):
            return float(obj)
        return super().default(obj)

import json
from datetime import date, timedelta
from decimal import Decimal

from django.contrib.auth.decorators import login_required
from django.db.models import Avg, Count, FloatField
from django.shortcuts import render, redirect
from django.utils import timezone

@login_required
def analytics_dashboard(request):
    if not request.user.is_superuser:
        return redirect('staff_dashboard')

    # Custom JSON encoder for Decimal values
    class DecimalEncoder(json.JSONEncoder):
        def default(self, obj):
            if isinstance(obj, Decimal):
                return float(obj)
            return super().default(obj)

    # 1. Class Performance (Average Score per Class)
    class_performance = list(
        ClassRoom.objects.annotate(
            avg_score=Avg('students__scores__score', output_field=FloatField())
        )
        .filter(avg_score__isnull=False)
        .values('id', 'name', 'avg_score')
        .order_by('-avg_score')
    )

    # Calculate School-Wide Average (from class averages)
    school_wide_avg = 0.0
    if class_performance:
        total = sum(item['avg_score'] or 0 for item in class_performance)
        school_wide_avg = round(total / len(class_performance), 1)

    # 2. Top 10 Students Overall
    top_students_list = list(
        Student.objects.annotate(
            avg_score=Avg('scores__score', output_field=FloatField()),
            total_scores=Count('scores')
        )
        .filter(total_scores__gte=3, avg_score__isnull=False)
        .order_by('-avg_score')[:10]
        .values(
            'user__first_name', 'user__last_name',
            'roll_number', 'classroom__name',
            'avg_score'
        )
    )

    top_performer_avg = None
    if top_students_list:
        top_performer_avg = round(top_students_list[0]['avg_score'], 1)

    # 3. Bottom 10 Students (Needs Attention)
    bottom_students = list(
        Student.objects.annotate(
            avg_score=Avg('scores__score', output_field=FloatField()),
            total_scores=Count('scores')
        )
        .filter(total_scores__gte=1, avg_score__isnull=False)
        .order_by('avg_score')[:10]
        .values(
            'user__first_name', 'user__last_name',
            'roll_number', 'classroom__name',
            'avg_score'
        )
    )

    # 4. Subject Performance Radar
    subjects_radar = list(
        Subject.objects.annotate(
            avg_score=Avg('scores__score', output_field=FloatField())
        )
        .filter(avg_score__isnull=False)
        .values('name', 'avg_score')
        .order_by('name')
    )

    # 5. 30-Day Attendance Trend
    end_date = timezone.localdate()
    start_date = end_date - timedelta(days=29)

    attendance_data = []
    total_present = 0
    total_absent = 0

    current = start_date
    while current <= end_date:
        present = Attendance.objects.filter(date=current, status='present').count()
        absent = Attendance.objects.filter(date=current, status='absent').count()

        total_present += present
        total_absent += absent

        attendance_data.append({
            'date': current.strftime('%b %d'),
            'day': current.strftime('%a'),
            'present': present,
            'absent': absent,
            'total': present + absent,
        })
        current += timedelta(days=1)

    # Attendance Rate
    total_attended = total_present + total_absent
    attendance_rate = round((total_present / total_attended * 100), 2) if total_attended > 0 else 0

    # Context
    context = {
        'class_performance': json.dumps(class_performance, cls=DecimalEncoder),
        'top_students': json.dumps(top_students_list, cls=DecimalEncoder),
        'bottom_students': json.dumps(bottom_students, cls=DecimalEncoder),
        'subjects_radar': json.dumps(subjects_radar, cls=DecimalEncoder),
        'attendance_data': json.dumps(attendance_data),
        'attendance_rate': attendance_rate,
        'total_present': total_present,
        'total_absent': total_absent,
        'school_wide_avg': school_wide_avg,
        'top_performer_avg': top_performer_avg,
    }

    return render(request, 'core/analytics_dashboard.html', context)

import json
from decimal import Decimal

from django.contrib.auth.decorators import login_required
from django.db.models import Avg, Count, FloatField
from django.shortcuts import render
from django.core.serializers.json import DjangoJSONEncoder

@login_required
def performance_dashboard(request):
    selected_grade = request.GET.get('grade')  # '' means All Grades

    # Base querysets
    students_qs = Student.objects.select_related('user', 'classroom')
    scores_qs = Score.objects.select_related('student', 'subject', 'student__user', 'student__classroom')

    selected_grade_name = "All Grades"
    selected_classroom = None

    if selected_grade:
        try:
            selected_classroom = ClassRoom.objects.get(id=selected_grade)
            selected_grade_name = selected_classroom.name

            students_qs = students_qs.filter(classroom=selected_classroom)
            scores_qs = scores_qs.filter(student__classroom=selected_classroom)
        except (ValueError, ClassRoom.DoesNotExist):
            selected_grade = ''
            selected_grade_name = "All Grades"

    # === BIG 4 STATS ===
    total_students = students_qs.count()

    score_stats = scores_qs.aggregate(
        avg_score=Avg('score', output_field=FloatField()),
        total_scores=Count('id')
    )
    avg_score = round(score_stats['avg_score'] or 0.0, 1)

    passing_scores = scores_qs.filter(score__gte=60).count()
    pass_rate = (
        round(passing_scores / score_stats['total_scores'] * 100, 1)
        if score_stats['total_scores'] > 0 else 0.0
    )

    at_risk = students_qs.annotate(
        student_avg=Avg('scores__score', output_field=FloatField())
    ).filter(student_avg__lt=60).count()

    # === SUBJECT AVERAGES FOR RADAR CHART ===
    subject_averages_qs = Subject.objects.annotate(
        avg=Avg('scores__score', output_field=FloatField())
    ).filter(avg__isnull=False)

    if selected_grade:
        subject_averages_qs = subject_averages_qs.filter(
            scores__student__classroom=selected_classroom
        ).distinct()

    subject_averages = list(
        subject_averages_qs.values('name', 'avg').order_by('name')
    )
    for item in subject_averages:
        item['avg'] = round(item['avg'] or 0, 1)

    # === TOP 10 & BOTTOM 10 STUDENTS ===
    ranked_students = students_qs.annotate(
        avg_score=Avg('scores__score', output_field=FloatField())
    ).filter(avg_score__isnull=False).values(
        'user__first_name', 'user__last_name',
        'roll_number', 'avg_score',
        'classroom__name'
    ).order_by('-avg_score')

    top_10 = []
    bottom_10 = []

    for s in ranked_students:
        avg = round(s['avg_score'] or 0, 1)
        student_data = {
            'name': f"{s['user__first_name']} {s['user__last_name']}".strip(),
            'roll_number': s['roll_number'] or '-',
            'classroom': s['classroom__name'],
            'avg': avg,
        }
        top_10.append(student_data)
        bottom_10.insert(0, student_data)  # Efficient reverse

    top_10 = top_10[:10]
    bottom_10 = bottom_10[:10]

    # === GENDER DISTRIBUTION PIE CHART ===
    gender_data = list(
        students_qs.values('user__gender')
        .annotate(count=Count('id'))
        .order_by('user__gender')
    )

    gender_pie = []
    gender_map = {'M': 'Male', 'F': 'Female', 'O': 'Other', None: 'Not Specified'}
    for item in gender_data:
        label = gender_map.get(item['user__gender'], 'Not Specified')
        gender_pie.append({'label': label, 'value': item['count']})

    # === STUDENT PERFORMANCE MATRIX ===
    subjects = list(Subject.objects.all().order_by('name'))

    relevant_scores = scores_qs.values(
        'student_id', 'subject__name', 'score'
    )

    scores_by_student = {}
    for sc in relevant_scores:
        if sc['score'] is not None:
            scores_by_student.setdefault(sc['student_id'], {})[sc['subject__name']] = float(sc['score'])

    student_matrix = []
    for student in students_qs:
        student_scores = scores_by_student.get(student.id, {})
        total = sum(student_scores.values())
        count = len(student_scores)
        average = round(total / count, 1) if count > 0 else 0.0

        student_matrix.append({
            'id': student.id,
            'name': student.user.get_full_name() or student.user.username,
            'roll_number': student.roll_number or '-',
            'classroom': student.classroom.name if student.classroom else 'N/A',
            'scores': student_scores,
            'total': round(total, 1),
            'average': average,
        })

    student_matrix.sort(key=lambda x: x['average'], reverse=True)
    student_matrix = student_matrix[:200]  # Safe limit

    # === CONTEXT ===
    context = {
        'total_students': total_students,
        'avg_score': avg_score,
        'pass_rate': pass_rate,
        'at_risk': at_risk,

        'classrooms': ClassRoom.objects.all().order_by('name'),
        'selected_grade': selected_grade,
        'selected_grade_name': selected_grade_name,

        'subject_averages': json.dumps(subject_averages, cls=DjangoJSONEncoder),
        'top_10': json.dumps(top_10, cls=DjangoJSONEncoder),
        'bottom_10': json.dumps(bottom_10, cls=DjangoJSONEncoder),
        'gender_pie': json.dumps(gender_pie, cls=DjangoJSONEncoder),

        'student_matrix': student_matrix,
        'subjects': subjects,
    }

    return render(request, 'core/performance_dashboard.html', context)

from django.contrib.auth import login
from django.contrib.auth.views import LoginView
from django.shortcuts import redirect
from django.contrib import messages
from django.urls import reverse

class CustomLoginView(LoginView):
    template_name = 'core/login.html'

    def form_valid(self, form):
        user = form.get_user()
        login(self.request, user)

        # Role-based redirect
        if user.is_superuser:
            return redirect('staff_dashboard')
        elif user.role == 'teacher':
            return redirect('teacher_dashboard')
        elif user.role == 'student':
            return redirect('student_dashboard')
        elif user.role == 'parent':
            return redirect('parent_dashboard')
        else:
            return redirect('teacher_dashboard')  # fallback

    def form_invalid(self, form):
        messages.error(self.request, "Invalid username or password. Please try again.")
        return super().form_invalid(form)